"""Generador de reporte Word automático con mapa embebido.

Uso:
    from agente.reporte import generar_reporte_operacion
    path = generar_reporte_operacion(ciudad="Medellín", fecha_inicio="2026-04-01", fecha_fin="2026-04-27")
"""
from __future__ import annotations

import io
import os
import sys
from datetime import datetime
from pathlib import Path

_ROOT = Path(__file__).resolve().parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agente.herramientas import consultar_metricas, generar_mapa, CIUDADES_NOMBRE
from agente.captura import capturar_mapa_html


# ── Helpers de formato ────────────────────────────────────────────────────────

def _color_celda(celda, hex_color: str):
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _pct_color(valor: float | None) -> str:
    """Verde/amarillo/rojo según valor porcentual."""
    if valor is None:
        return "F5F5F5"
    if valor >= 60:
        return "C8E6C9"
    if valor >= 35:
        return "FFF9C4"
    return "FFCDD2"


# ── Función principal ─────────────────────────────────────────────────────────

def generar_reporte_operacion(
    ciudad: str,
    fecha_inicio: str | None = None,
    fecha_fin: str | None = None,
    output_dir: str = "Claude/outputs",
    capturar: bool = True,
) -> str:
    """Genera reporte Word de operación para una ciudad.

    Args:
        ciudad:       Nombre de ciudad.
        fecha_inicio: YYYY-MM-DD (default: primer día del mes).
        fecha_fin:    YYYY-MM-DD (default: hoy).
        output_dir:   Carpeta donde guardar el .docx.
        capturar:     Si True, genera screenshot del mapa e incluye imagen.

    Returns:
        Ruta al archivo .docx generado.
    """
    from agente.herramientas import _resolver_ciudad, _fechas_por_defecto

    if fecha_inicio is None or fecha_fin is None:
        fecha_inicio, fecha_fin = _fechas_por_defecto()

    id_centroope, nombre_ciudad = _resolver_ciudad(ciudad)

    # ── 1. Obtener datos ──────────────────────────────────────────────────────
    datos = consultar_metricas(ciudad, fecha_inicio, fecha_fin)
    resumen = datos.get("resumen", {})
    promotores = datos.get("promotores", [])

    # ── 2. Generar mapa y captura ─────────────────────────────────────────────
    png_path = None
    if capturar:
        try:
            mapa_res = generar_mapa(ciudad, fecha_inicio, fecha_fin)
            html_path = mapa_res.get("html_path", "")
            if html_path and os.path.exists(html_path):
                png_path = capturar_mapa_html(html_path, delay_ms=3000)
        except Exception as e:
            print(f"[reporte] Advertencia — no se pudo capturar mapa: {e}")

    # ── 3. Construir documento ────────────────────────────────────────────────
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # Portada
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f"Reporte de Operación de Campo")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{nombre_ciudad}  ·  {fecha_inicio} → {fecha_fin}").font.size = Pt(13)

    doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}").runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Resumen ejecutivo
    doc.add_heading("Resumen ejecutivo", level=1)
    items = [
        ("Promotores activos",   resumen.get("n_promotores", "—")),
        ("Clientes visitados",   f"{resumen.get('total_clientes', 0):,}"),
        ("% Contactabilidad",    f"{resumen.get('prom_contactabilidad_pct', '—')}%"),
        ("% Captación",          f"{resumen.get('prom_captacion_pct', '—')}%"),
        ("% Conversión",         f"{resumen.get('prom_conversion_pct', '—')}%"),
    ]
    for label, val in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(str(val))

    doc.add_paragraph()

    # Imagen del mapa
    if png_path and os.path.exists(png_path):
        doc.add_heading("Mapa de cobertura", level=1)
        doc.add_picture(png_path, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # Tabla de métricas
    doc.add_heading("Métricas por promotor", level=1)

    if promotores:
        cols = ["nombre_promotor", "clientes_total", "pct_contactabilidad",
                "pct_nofiel_contactable", "pct_conversion"]
        headers = ["Promotor", "Clientes", "% Contactab.", "% Captación", "% Conversión"]

        tabla = doc.add_table(rows=1, cols=len(headers))
        tabla.style = "Table Grid"

        # Encabezado
        for i, h in enumerate(headers):
            cell = tabla.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            _color_celda(cell, "1A1A2E")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Filas
        for p in promotores:
            row = tabla.add_row().cells
            row[0].text = str(p.get("nombre_promotor", ""))
            row[1].text = str(p.get("clientes_total", ""))

            for j, col in enumerate(["pct_contactabilidad", "pct_nofiel_contactable", "pct_conversion"]):
                val = p.get(col)
                row[j + 2].text = f"{val:.1f}%" if val is not None else "—"
                _color_celda(row[j + 2], _pct_color(val))

    # ── 4. Guardar ────────────────────────────────────────────────────────────
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    fecha_tag = fecha_fin.replace("-", "")
    filename = f"Reporte_{nombre_ciudad}_{fecha_tag}.docx"
    out_path = str(Path(output_dir) / filename)
    doc.save(out_path)

    print(f"[reporte] Guardado: {out_path}")
    return out_path
